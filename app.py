import streamlit as st
import google.generativeai as genai
from pdf2image import convert_from_bytes
from PIL import Image
import io
import base64
import time
import os
from typing import List, Optional, Tuple
import zipfile
from docx import Document
import tempfile
import logging
from dataclasses import dataclass
from enum import Enum
import re

# --- Configuration ---
# Configure logging to get detailed output in the console
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Enums and Data Classes ---

class Language(Enum):
    """Enumeration for supported document languages."""
    ENGLISH = "English"
    NEPALI = "Nepali"

class OutputFormat(Enum):
    """Enumeration for supported output file formats."""
    MARKDOWN = "md"
    TEXT = "txt"
    DOCX = "docx"

@dataclass
class ProcessingStatus:
    """A data class to hold the status of the document processing."""
    total_pages: int = 0
    processed_pages: int = 0
    failed_pages: int = 0
    current_page: int = 0
    status_message: str = "Initializing..."
    
    @property
    def progress_percentage(self) -> float:
        """Calculate the progress percentage."""
        if self.total_pages == 0:
            return 0.0
        return (self.processed_pages / self.total_pages) * 100

# --- Core Logic: DocumentProcessor Class ---

class DocumentProcessor:
    """Handles the core logic of converting documents to Markdown."""
    def __init__(self, api_key: str):
        """Initialize the document processor with the Gemini API key."""
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash-latest') # Using the latest flash model
        
    def get_language_prompt(self, language: Language, user_prompt: str = "") -> str:
        """Get the language-specific instruction prompt for the Gemini API, with optional user instructions."""
        base_prompt = f"""
You are an expert document digitization assistant. Your task is to convert this scanned document image into clean, properly structured Markdown format.

CRITICAL REQUIREMENTS:
1.  **Preserve Structure**: Maintain all headings (H1, H2, H3), paragraphs, lists (bulleted and numbered), and tables.
2.  **Clean Output**: Remove scanning artifacts, noise, and page skew. Do not invent content.
3.  **Accurate Text**: Ensure the highest possible OCR accuracy.
4.  **Markdown Formatting**: Use standard Markdown syntax.
5.  **Language**: The document is in **{language.value}**. Process it with careful attention to the correct script, grammar, and spelling. For Nepali, ensure proper Devanagari script handling.
6.  **No Extra Text**: Return ONLY the Markdown content of the page. Do not add any explanations, apologies, or summaries like "Here is the markdown".
"""
        if user_prompt and user_prompt.strip():
            base_prompt += f"\n\nADDITIONAL USER INSTRUCTIONS:\n{user_prompt.strip()}\n"
        return base_prompt
    
    def pdf_to_images(self, pdf_bytes: bytes) -> List[Image.Image]:
        """Convert PDF bytes into a list of PIL Images using system Poppler."""
        try:
            images = convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"Successfully converted PDF to {len(images)} images.")
            return images
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            st.error(f"PDF conversion failed. The file might be corrupted or password-protected. Error: {e}")
            raise

    def process_image_with_gemini(self, image: Image.Image, language: Language, user_prompt: str = "") -> Optional[str]:
        """Process a single image using the Gemini API and return Markdown text, including user prompt."""
        try:
            prompt = self.get_language_prompt(language, user_prompt)
            response = self.model.generate_content([prompt, image])
            
            # Validate response before returning
            if response.text:
                return response.text.strip()
            else:
                logger.warning("Empty response from Gemini API for a page.")
                return "<!-- This page could not be processed. -->"
                
        except Exception as e:
            logger.error(f"Error processing image with Gemini: {e}")
            # Provide a specific error message for API issues
            if "API key not valid" in str(e):
                st.error("Authentication Error: Your Gemini API key is not valid. Please check and re-enter it.")
            return f"<!-- Page processing failed due to an API error: {e} -->"

    def process_document(self, file_bytes: bytes, file_type: str, language: Language, 
                        status_callback=None, user_prompt: str = "") -> Tuple[str, ProcessingStatus]:
        """
        Process an entire document (PDF or image).
        This function orchestrates the conversion, page by page, and uses a callback
        to update the Streamlit UI in real-time.
        """
        status = ProcessingStatus()
        markdown_pages = []
        
        try:
            images = self.pdf_to_images(file_bytes) if file_type == "pdf" else [Image.open(io.BytesIO(file_bytes))]
            status.total_pages = len(images)
            
            if status_callback:
                status_callback(status, markdown_pages)

            for i, image in enumerate(images):
                status.current_page = i + 1
                status.status_message = f"Analyzing page {status.current_page} of {status.total_pages}..."
                if status_callback:
                    status_callback(status, markdown_pages)
                
                page_content = self.process_image_with_gemini(image, language, user_prompt)
                
                if page_content:
                    markdown_pages.append(page_content)
                    status.processed_pages += 1
                    logger.info(f"Successfully processed page {i + 1}")
                else:
                    status.failed_pages += 1
                    logger.warning(f"Failed to process page {i + 1}")
                    markdown_pages.append(f"<!-- Page {i + 1}: Processing failed -->")
                
                status.status_message = f"Page {status.current_page} complete."
                if status_callback:
                    status_callback(status, markdown_pages)
                
                time.sleep(1) # Small delay to respect API rate limits and allow UI to update smoothly

            status.status_message = "Processing complete!"
            if status_callback:
                status_callback(status, markdown_pages)
            
            return "\n\n---\n\n".join(markdown_pages), status
            
        except Exception as e:
            logger.error(f"Critical error in document processing: {e}")
            status.status_message = f"Error: {e}"
            if status_callback:
                status_callback(status, markdown_pages)
            raise

# --- File Conversion Utilities ---

def convert_to_format(markdown_content: str, output_format: OutputFormat) -> bytes:
    """Convert the final Markdown content to the specified output format."""
    if output_format == OutputFormat.MARKDOWN:
        return markdown_content.encode('utf-8')
    
    elif output_format == OutputFormat.TEXT:
        # Improved regex for a cleaner text conversion
        text = re.sub(r'#+\s*', '', markdown_content)      # Remove headers
        text = re.sub(r'[\*_`]', '', text)                 # Remove bold/italic/code markers
        text = re.sub(r'---\s*', '\n', text)               # Replace separators with a newline
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)    # Keep link text, remove URL
        return text.encode('utf-8')
    
    elif output_format == OutputFormat.DOCX:
        doc = Document()
        # Correctly split the content into lines for processing
        for line in markdown_content.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Handle headings
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            # Handle bullet points (both * and -)
            elif line.startswith(('* ', '- ')):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()

# --- Streamlit UI ---

def main():
    """Main function to run the Streamlit application."""
    st.set_page_config(page_title="DocuMint AI", page_icon="📄", layout="wide")

    # --- Session State for API Key ---
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = ''

    # --- Sidebar ---
    with st.sidebar:
        st.markdown("""
            <div style="display: flex; align-items: center; gap: 10px;">
                <img src="https://img.icons8.com/ios-filled/50/000000/document--v1.png" width="32"/>
                <h2 style="margin-bottom: 0;">DocuMint AI</h2>
            </div>
        """, unsafe_allow_html=True)
        st.markdown("<span style='color: #6c757d;'>AI-powered document conversion to Markdown, Text, and DOCX.</span>", unsafe_allow_html=True)
        st.markdown("---")
        st.header("⚙️ Configuration")

        api_key_input = st.text_input(
            "Gemini API Key",
            type="password",
            value=st.session_state['api_key'],
            help="Get your key from Google AI Studio."
        )
        if api_key_input:
            st.session_state['api_key'] = api_key_input
        api_key = st.session_state['api_key']

        language = st.selectbox(
            "Document Language",
            options=[lang for lang in Language],
            format_func=lambda x: x.value
        )
        output_formats = st.multiselect(
            "Output Formats",
            options=[fmt for fmt in OutputFormat],
            default=[OutputFormat.MARKDOWN, OutputFormat.DOCX],
            format_func=lambda x: f".{x.value}"
        )

        st.info("Built by Chhatramani Yadav using Streamlit & Gemini.")

        st.markdown("### 📋 Instructions")
        st.markdown("""
        1. Enter your Gemini API key
        2. Select document language
        3. Upload PDF or image files
        4. Choose output formats
        5. Click 'Start Process'
        """)

    # --- Main Page Layout ---
    st.markdown("""
        <div style="background: linear-gradient(90deg, #4e54c8 0%, #8f94fb 100%); padding: 2rem 2rem 1rem 2rem; border-radius: 12px; margin-bottom: 2rem; color: white;">
            <h1 style="margin-bottom: 0.2em;">📄 DocuMint AI Dashboard</h1>
            <p style="font-size: 1.2em; margin-bottom: 0;">Upload your PDFs or images and watch the AI convert them in real-time.</p>
        </div>
    """, unsafe_allow_html=True)


    # --- Additional Prompt Box ---
    st.markdown("<h4 style='margin-top:0;'>Additional Instructions (Optional)</h4>", unsafe_allow_html=True)
    user_prompt = st.text_area(
        "Add any extra instructions for the AI (e.g., focus on tables, ignore watermarks, etc.)",
        value="",
        height=80,
        key="user_prompt_box"
    )

    uploaded_files = st.file_uploader(
        "Choose PDF or image files",
        type=['pdf', 'png', 'jpg', 'jpeg'],
        accept_multiple_files=True,
        help="You can upload multiple files at once."
    )

    # --- Start Processing Button and Notices ---
    process_button = st.button(
        "🚀 Start Processing",
        disabled=not (api_key and uploaded_files),
        type="primary"
    )
    if not api_key:
        st.warning("⚠️ Please enter your Gemini API key in the sidebar.")
    elif not uploaded_files:
        st.info("ℹ️ Upload one or more files to get started.")

    # --- Processing Logic ---
    if process_button and api_key and uploaded_files:
        try:
            processor = DocumentProcessor(api_key)
            all_results = {}

            # Create a container for all results
            results_area = st.container()
            results_area.header(":sparkles: Conversion Results")

            for file_idx, uploaded_file in enumerate(uploaded_files):
                with results_area.expander(f"**Processing: {uploaded_file.name}**", expanded=True):

                    # Create placeholders for dynamic content. These will be updated live.
                    status_placeholder = st.empty()
                    live_preview_placeholder = st.empty()

                    def update_status_ui(status: ProcessingStatus, current_md_pages: List[str]):
                        """
                        This is the callback function. It updates the UI elements in-place.
                        """
                        # --- Status Metrics Update (Single Line) ---
                        with status_placeholder.container():
                            cols = st.columns([2, 1, 1, 1])
                            cols[0].markdown(f"<span style='font-size:1.1em;'><b>Status:</b> {status.status_message}</span>", unsafe_allow_html=True)
                            cols[1].metric("Total Pages", status.total_pages)
                            cols[2].metric("Processed", f"{status.processed_pages}/{status.total_pages}")
                            cols[3].metric("Failed", status.failed_pages)
                            if status.total_pages > 0:
                                st.progress(status.progress_percentage / 100)

                        # --- Live Markdown Preview Update ---
                        with live_preview_placeholder.container():
                            st.markdown("<h4 style='margin-bottom:0.5em;'>Live Markdown Preview</h4>", unsafe_allow_html=True)
                            # Join the pages processed so far for a live view
                            live_content = "\n\n---\n\n".join(current_md_pages)
                            st.markdown(live_content, unsafe_allow_html=True)

                    # Start the document processing
                    file_bytes = uploaded_file.read()
                    file_type = "pdf" if uploaded_file.type == "application/pdf" else "image"
                    markdown_content, final_status = processor.process_document(
                        file_bytes, file_type, language, update_status_ui, user_prompt
                    )

                    # Store final results
                    all_results[uploaded_file.name] = markdown_content

                    # --- Final Editable Area and Download Buttons ---
                    st.success(f"✅ <b>{uploaded_file.name}</b> processed successfully!", icon="✅")
                    st.markdown("<h4>Edit & Download</h4>", unsafe_allow_html=True)
                    edited_content = st.text_area(
                        "You can make final edits to the Markdown here:",
                        value=markdown_content,
                        height=300,
                        key=f"edit_{file_idx}"
                    )
                    # Update results with any edits made by the user
                    all_results[uploaded_file.name] = edited_content

                    download_cols = st.columns(len(output_formats))
                    for i, fmt in enumerate(output_formats):
                        with download_cols[i]:
                            converted_bytes = convert_to_format(edited_content, fmt)
                            st.download_button(
                                label=f"⬇️ Download .{fmt.value}",
                                data=converted_bytes,
                                file_name=f"{os.path.splitext(uploaded_file.name)[0]}.{fmt.value}",
                                mime=f"text/{fmt.value}" if fmt != OutputFormat.DOCX else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )

            # --- Bulk Download Option ---
            if len(uploaded_files) > 1:
                with st.sidebar:
                    st.header("📦 Bulk Download")
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for filename, content in all_results.items():
                            base_name = os.path.splitext(filename)[0]
                            for fmt in output_formats:
                                converted_bytes = convert_to_format(content, fmt)
                                zf.writestr(f"{base_name}/{base_name}.{fmt.value}", converted_bytes)

                    st.download_button(
                        label="📥 Download All (ZIP)",
                        data=zip_buffer.getvalue(),
                        file_name="documint_ai_converted_files.zip",
                        mime="application/zip",
                        use_container_width=True
                    )

        except Exception as e:
            st.error(f"❌ A critical error occurred during processing: {e}")
            logger.error(f"Top-level processing error: {e}", exc_info=True)

if __name__ == "__main__":
    main()
